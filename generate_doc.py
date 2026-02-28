from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

doc = Document()

# Set RTL direction for the document
# Note: python-docx has limited built-in RTL support, but we can align text to right 
# and use Arabic text.

# Add Title
title = doc.add_heading('دليل المستخدم: دورة حياة المريض في القسم الداخلي (التنويم والمحاسبة)', level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph()

def add_step(title, description, details):
    p_title = doc.add_heading(title, level=2)
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    p_desc = doc.add_paragraph(description)
    p_desc.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    for detail in details:
        p_detail = doc.add_paragraph(style='List Bullet')
        p_detail.text = detail
        p_detail.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
    doc.add_paragraph()

add_step(
    "1. إضافة المريض وحجزه المبدئي",
    "تبدأ دورة المريض من مكتب الاستقبال لتسجيل بياناته الأساسية وفتح تذكرة.",
    [
        "يتم البحث عن المريض بالاسم، رقم الملف، أو الهوية للتأكد من عدم ازدواجية التسجيل.",
        "في حال المريض الجديد، يتم إدخال بياناته الأساسية (الاسم، الجوال، الجنسية، إلخ).",
        "يمر المريض بالعيادة أو الطوارئ، حيث يقرر الطبيب تقييمه ويطلب تنويمه في القسم الداخلي."
    ]
)

add_step(
    "2. إجراءات التنويم في القسم الداخلي (Admission)",
    "بعد قرار الطبيب بالتنويم، يقوم قسم تقارير الدخول أو الاستقبال بحجز غرفته.",
    [
        "يتم تحديد تاريخ متوقع للخروج واسم الطبيب المعالج.",
        "يتم اختيار الغرفة والسرير المناسبين من قائمة الغرف المتاحة.",
        "يضاف مبلغ 'تأمين مقدّم' أو 'دفعة تحت حساب التنويم' بناءً على سياسة المستشفى.",
        "تتغير حالة السرير إلى 'مشغول'، ويفتح سجل مالي مفصل للمريض المنوم."
    ]
)

add_step(
    "3. طلب الخدمات والتحاليل والمستهلكات (Medical Orders)",
    "أثناء إقامة المريض، يطلب الأطباء والممرضون خدمات متعددة تسجل جميعها على حسابه بشكل آلي.",
    [
        "طلب التحاليل الطبية (المختبر) والأشعة: يتم عبر النظام وترسل مباشرة للقسم المختص لجدولتها.",
        "طلب المستهلكات الطبية والأدوية: تضاف بكمياتها ويتم خصمها من المخزون فور صرفها للمريض.",
        "أجور إقامة الغرف: يتم احتسابها تلقائياً بشكل يومي، وإضافتها لفاتورة المريض المجمّعة.",
        "تسجل هذه المطالبات بقيمتها في حساب المريض كـ 'غير مسددة' حتى انتهاء التنويم."
    ]
)

add_step(
    "4. تسجيل العمليات الجراحية (Surgical Operations)",
    "إذا استدعت حالة المريض إجراء عملية جراحية، يتم تحويله للعمليات وتسجيل تفاصيلها.",
    [
        "تضاف العملية لملف المريض، بما في ذلك بيانات الجرّاح وطبيب التخدير ونوع العملية.",
        "يتم توزيع أتعاب الأطباء وحصة المستشفى آلياً في النظام.",
        "تُدرج تكلفة العملية في إجمالي حساب إقامة المريض المنوم (الفاتورة المجمعة العظمى)."
    ]
)

add_step(
    "5. الخروج الطبي والنهائي (Medical Discharge)",
    "عند تحسن حالة المريض وتصريح الطبيب بخروجه، تبدأ إجراءات الخروج الطبية لتنتقل للإجراءات المالية.",
    [
        "يقوم الطبيب بكتابة 'ملاحظات الخروج' وإصدار الإذن الطبي.",
        "تقوم الممرضة بتأكيد خروج المريض من الغرفة، لتبدأ عملية تنظيف السرير (Cleaning).",
        "تتوقف عملية احتساب أيام الإقامة آلياً."
    ]
)

add_step(
    "6. المحاسبة وإصدار الفاتورة النهائية للمنوم (Billing & Invoicing)",
    "يتوجه المريض أو مرافقه لقسم المحاسبة لتسوية الملف المالي والحصول على فاتورته الشاملة.",
    [
        "يقوم المحاسب بتوليد 'فاتورة التنويم المجمعة' التي تتضمن: رسوم الإقامة، التحاليل، الأشعة، العمليات، والأدوية.",
        "تُخصم أي دفعات مقدمة (Advance Payments) دفعها المريض سابقاً تلقائياً من إجمالي الفاتورة.",
        "يتم احتساب الضرائب ونسب خصم التأمين أو المؤسسات لتحديد 'المبلغ المتبقي الصافي'."
    ]
)

add_step(
    "7. دفع المبالغ المستحقة والتسوية (Payments & Settlement)",
    "وهي الخطوة الأخيرة لإنهاء ملف التنويم بشكل رسمي وإتمام القيود المحاسبية.",
    [
        "يقوم المريض بدفع الرصيد المتبقي (نقدي، بطاقة، إلخ).",
        "يصدر النظام 'سند قبض' بشكل آلي ويقيد المبالغ في حساب الصندوق/البنك وحسابات الإيرادات.",
        "تصبح حالة الفاتورة 'مسددة بالكامل'، ويسمح للمريض بمغادرة المستشفى وبيده 'تقرير خدمات المريض' مفصلاً ومدفوعاً."
    ]
)

# Save to Desktop
import winreg
def get_desktop_path():
    try:
        key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, r'Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders')
        return winreg.QueryValueEx(key, 'Desktop')[0]
    except:
        return os.path.expanduser('~\\Desktop')

desktop_path = get_desktop_path()
file_path = os.path.join(desktop_path, 'دليل_التنويم_والمحاسبة.docx')
doc.save(file_path)

print(f"Document successfully created at: {file_path}")
